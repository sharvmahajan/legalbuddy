import streamlit as st
import os
import tempfile
import time
import pandas as pd
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from langchain.document_loaders import PyPDFLoader, UnstructuredWordDocumentLoader, TextLoader, CSVLoader, UnstructuredExcelLoader
from langchain.schema import Document
from docx import Document as DocxDocument
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import json

# Load environment variables
load_dotenv()

# Custom CSS for the cards
st.markdown("""
<style>
    div.stButton > button {
        width: 100%;
        height: 250px;  /* Increased height */
        background-color: #f0f2f6;
        border: none;
        border-radius: 10px;
        padding: 20px;
        margin: 10px 0;
        transition: transform 0.3s ease, box-shadow 0.3s ease;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        text-align: center;
        white-space: normal !important;  /* Allow text wrapping */
        word-wrap: break-word;
    }
    div.stButton > button:hover {
        transform: translateY(-5px);
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        background-color: #e0e2e6;
    }
    div.stButton > button p {
        font-size: 20px;  /* Slightly reduced font size */
        font-weight: bold;
        margin-bottom: 15px;
        color: #0f1629;
        line-height: 1.4;
    }
    div.stButton > button small {
        font-size: 14px;
        color: #485164;
        line-height: 1.4;
    }
</style>
""", unsafe_allow_html=True)

st.title("📚 LegalBuddy")
st.markdown("### Your AI-powered Legal Document Analyzer")

# Initialize session state variables
if 'selected_option' not in st.session_state:
    st.session_state.selected_option = None
if 'vectors' not in st.session_state:
    st.session_state.vectors = None
if 'processed_files' not in st.session_state:
    st.session_state.processed_files = []
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# Load API keys from environment variables
groq_api_key = os.getenv('GROQ_API_KEY')
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")

# --------------------------- ENHANCED SYSTEM PROMPT ---------------------------
LEGAL_EXTRACTION_PROMPT = """You are an expert legal document analyzer. Your task is to meticulously extract and categorize specific legal information from documents.

EXTRACTION GUIDELINES:
1. Entities and Contact Details
   - Identify all parties involved
   - Extract full legal names
   - Capture addresses and contact information

2. Contract Timeline
   - Locate and extract start date (effective date of term sheet)
   - Identify end date or contract duration (expiration date of term sheet)
   - Note any key milestone dates (any dates relevent to the term sheet)

3. Scope of Agreement
   - Clearly define the purpose and scope of the document
   - Highlight key deliverables or services

4. Service Level Agreement (SLA)
   - Extract specific performance metrics
   - Identify response times and service standards

5. Penalty Clauses
   - Identify specific penalty conditions
   - Extract monetary or legal consequences for non-compliance, also state what counts as non-compliance

6. Confidentiality Provisions
   - Detail confidentiality obligations
   - Extract duration and scope of confidentiality

7. Renewal and Termination
   - Extract conditions for contract renewal (when contract ends)
   - Identify termination clauses
   - Note notice periods

8. Commercial Terms
   - Extract payment terms
   - Identify pricing structures
   - Note invoicing and payment schedules

9. Risk and Assumptions
   - Identify and list potential risks
   - Extract any mitigation strategies or assumptions

If any section is not present, provide a clear statement about the absence of that specific information.

RESPONSE FORMAT:
- Use clear, structured headings
- Provide specific details where available
- Clearly indicate if information is missing
"""

# Prompts for different extraction tasks
extraction_prompt = ChatPromptTemplate.from_template(
    f"""
{LEGAL_EXTRACTION_PROMPT}

📜 **Document Context**:
<context>
{{context}}
</context>

🔍 **Extraction Task**: Extract and categorize all available legal information from the document.
"""
)

qa_prompt = ChatPromptTemplate.from_template(
    f"""
You are a legal document assistant. Provide precise and contextual answers.

📜 **Document Context**:
<context>
{{context}}
</context>

🔍 **User Question**: {{input}}

Provide a clear, concise answer based strictly on the document context.
"""
)

# --------------------------- DOCUMENT PROCESSOR ---------------------------
def process_uploaded_file(uploaded_file):
    file_name = uploaded_file.name
    file_extension = file_name.split('.')[-1].lower()

    # Create a temporary file to store the uploaded file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}') as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name

    # Process different file types
    try:
        if file_extension == 'pdf':
            loader = PyPDFLoader(tmp_path)
            documents = loader.load()
        elif file_extension == 'txt':
            loader = TextLoader(tmp_path)
            documents = loader.load()
        elif file_extension == 'docx':
            documents = []
            doc = DocxDocument(tmp_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    documents.append(Document(page_content=para.text))
        elif file_extension == 'csv':
            loader = CSVLoader(tmp_path)
            documents = loader.load()
        elif file_extension == 'xlsx':
            loader = UnstructuredExcelLoader(tmp_path)
            documents = loader.load()
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return []
    finally:
        print("thik he")
        # Clean up the temporary file
        # os.unlink(tmp_path)

    return documents

def vector_embedding(documents):
    with st.spinner("Processing documents..."):
        st.session_state.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        st.session_state.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        st.session_state.final_documents = st.session_state.text_splitter.split_documents(documents)
        st.session_state.vectors = FAISS.from_documents(st.session_state.final_documents, st.session_state.embeddings)
    st.success("Documents processed successfully!")

def convert_to_json(extraction_text):
    """Convert the extracted text into a structured JSON format"""
    sections = {
        "entities_and_contacts": {},
        "contract_timeline": {},
        "scope": "",
        "sla_clauses": [],
        "penalty_clauses": [],
        "confidentiality": {},
        "renewal_termination": {},
        "commercial_terms": {},
        "risks_assumptions": []
    }
    
    # Parse the extraction text and populate sections
    current_section = None
    for line in extraction_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if any(key.replace('_', ' ').upper() in line.upper() for key in sections.keys()):
            current_section = next(key for key in sections.keys() if key.replace('_', ' ').upper() in line.upper())
            continue
            
        if current_section and line:
            if isinstance(sections[current_section], dict):
                # Split on first colon for key-value pairs
                if ':' in line:
                    key, value = line.split(':', 1)
                    sections[current_section][key.strip()] = value.strip()
            elif isinstance(sections[current_section], list):
                sections[current_section].append(line)
            else:
                sections[current_section] = line
    
    return sections
# --------------------------- INITIALIZE LLM ---------------------------
llm = ChatGroq(groq_api_key=groq_api_key, model_name="Llama3-8b-8192")

# --------------------------- MAIN APP LAYOUT ---------------------------
st.markdown("## 👋 Hello! What do you want to know today?")

# File Uploader
uploaded_files = st.file_uploader(
    "Upload your legal documents (PDF, DOCX, TXT, CSV, XLSX, etc.)", 
    type=['pdf', 'txt', 'docx', 'csv', 'xlsx'], 
    accept_multiple_files=True
)

if uploaded_files:
    new_files = [f.name for f in uploaded_files if f.name not in st.session_state.processed_files]
    if new_files:
        st.write("New files to process:", new_files)
        if st.button("🚀 Process Documents", use_container_width=True):
            all_documents = []
            for uploaded_file in uploaded_files:
                documents = process_uploaded_file(uploaded_file)
                all_documents.extend(documents)
            if all_documents:
                vector_embedding(all_documents)
                st.session_state.processed_files = [f.name for f in uploaded_files]

# Functionality Options
if st.session_state.vectors is not None:
    col1, col2, col3 = st.columns(3)

    # Card 1 - Document Extraction
    with col1:
        if st.button("🔍 Extract Key Details\n\nAutomatic extraction of critical legal information"):
            st.session_state.selected_option = 'extraction'

    # Card 2 - Chat with Documents
    with col2:
        if st.button("💬 Chat with Docs\n\nAsk specific questions about your documents"):
            st.session_state.selected_option = 'qa'

    # Card 3 - Maintain Chat History
    with col3:
        if st.button("📜 Chat History\n\nReview previous conversations"):
            st.session_state.selected_option = 'history'

    # --------------------------- MAIN FUNCTIONALITY ---------------------------
    if st.session_state.selected_option == 'extraction':
        with st.spinner("Extracting key details..."):
            document_chain = create_stuff_documents_chain(llm, extraction_prompt)
            retriever = st.session_state.vectors.as_retriever()
            retrieval_chain = create_retrieval_chain(retriever, document_chain)
            
            start = time.process_time()
            response = retrieval_chain.invoke({'input': 'Extract all key legal information from the document'})
            elapsed = time.process_time() - start
            
            # Convert to JSON
            json_data = convert_to_json(response['answer'])
            
            st.success(f"Analysis completed in {elapsed:.2f} seconds!")
            st.subheader("📑 Legal Document Insights")
            st.write(response['answer'])
            
            # Add download buttons
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 Download JSON",
                    data=json.dumps(json_data, indent=2),
                    file_name="legal_extraction.json",
                    mime="application/json"
                )
            with col2:
                st.download_button(
                    label="📥 Download Text",
                    data=response['answer'],
                    file_name="legal_extraction.txt",
                    mime="text/plain"
                )
            
            with st.expander("📚 Source Document Sections"):
                for i, doc in enumerate(response["context"]):
                    st.markdown(f"*Section {i+1}*")
                    st.write(doc.page_content)
                    st.divider()

    elif st.session_state.selected_option == 'qa':
        prompt1 = st.text_input("🔍 Ask a specific question about your documents")
        if prompt1:
            # Store user question in chat history
            st.session_state.chat_history.append(f"User: {prompt1}")
            
            with st.spinner("Finding answer..."):
                document_chain = create_stuff_documents_chain(llm, qa_prompt)
                retriever = st.session_state.vectors.as_retriever()
                retrieval_chain = create_retrieval_chain(retriever, document_chain)
                
                start = time.process_time()
                response = retrieval_chain.invoke({'input': prompt1})
                elapsed = time.process_time() - start
                
                # Store AI response in chat history
                st.session_state.chat_history.append(f"AI: {response['answer']}")
                
                st.success(f"Answer found in {elapsed:.2f} seconds!")
                st.subheader("💡 Answer")
                st.write(response['answer'])
                
                with st.expander("📚 Relevant Document Sections"):
                    for i, doc in enumerate(response["context"]):
                        st.markdown(f"*Section {i+1}*")
                        st.write(doc.page_content)
                        st.divider()

    elif st.session_state.selected_option == 'history':
        st.subheader("📜 Chat History")
        if st.session_state.chat_history:
            for message in st.session_state.chat_history:
                st.write(message)
            
            if st.button("Clear Chat History"):
                st.session_state.chat_history = []
                st.experimental_rerun()
        else:
            st.write("No chat history available.")

# --------------------------- FOOTER ---------------------------
st.markdown("---")
st.markdown("Made with ❤ by LegalBuddy Team")